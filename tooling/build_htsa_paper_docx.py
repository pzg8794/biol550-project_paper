#!/usr/bin/env python3
from __future__ import annotations

import re
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

PAPER_DIR = Path(__file__).resolve().parent.parent
TEX_PATH = PAPER_DIR / "HTSA_Paper.tex"
DOCX_PATH = PAPER_DIR / "HTSA_Paper.docx"

MEMBER_FILLS = {
    "nikhipara": "D9D2E9",
    "sampara": "B6D7A8",
    "piterpara": "F4CCCC",
}

STAGE_FILLS = {
    "Collection": "DDEBF7",
    "Cleaning": "E2F0D9",
    "Preparation": "FCE4D6",
    "Analysis": "E4DFEC",
}

CITATION_MAP = {
    "halawani2023ahr": "Halawani et al., 2023",
    "sraStudySRP618841": "NCBI SRA, 2026",
    "bioprojectPRJNA1322439": "NCBI BioProject, 2026",
    "geoGSE243308": "NCBI GEO, 2026",
    "crispdm2000": "Wirth & Hipp, 2000",
    "fastqc": "Andrews, 2010",
    "multiqc": "Ewels et al., 2016",
    "fastp": "Chen et al., 2018",
    "star": "Dobin et al., 2013",
    "deseq2": "Love et al., 2014",
    "gprofiler2021": "Kolberg et al., 2020",
    "geneOntology2021": "Gene Ontology Consortium, 2021",
    "kegg2025": "Kanehisa et al., 2025",
    "reactome2022": "Gillespie et al., 2022",
    "ensemblRelease115": "Ensembl, 2025",
    "sratoolkit": "NCBI, 2026",
}

REFERENCE_TEXT = [
    "Andrews, S. (2010). FastQC: A quality control tool for high throughput sequence data.",
    "Chen, S., Zhou, Y., Chen, Y., & Gu, J. (2018). fastp: An ultra-fast all-in-one FASTQ preprocessor. Bioinformatics, 34(17), i884-i890. https://doi.org/10.1093/bioinformatics/bty560",
    "Dobin, A., Davis, C. A., Schlesinger, F., Drenkow, J., Zaleski, C., Jha, S., Batut, P., Chaisson, M., & Gingeras, T. R. (2013). STAR: Ultrafast universal RNA-seq aligner. Bioinformatics, 29(1), 15-21. https://doi.org/10.1093/bioinformatics/bts635",
    "Ewels, P., Magnusson, M., Lundin, S., & Kaller, M. (2016). MultiQC: Summarize analysis results for multiple tools and samples in a single report. Bioinformatics, 32(19), 3047-3048. https://doi.org/10.1093/bioinformatics/btw354",
    "Gillespie, M., Jassal, B., Stephan, R., et al. (2022). The Reactome pathway knowledgebase 2022. Nucleic Acids Research, 50(D1), D687-D692. https://doi.org/10.1093/nar/gkab1028",
    "Halawani, D., Wang, Y., Li, J., et al. (2023). AhR restricts axon regeneration by balancing neuronal stress and growth response after injury. bioRxiv. https://doi.org/10.1101/2023.11.04.565649",
    "Kanehisa, M., Furumichi, M., Sato, Y., Kawashima, M., & Ishiguro-Watanabe, M. (2025). KEGG for taxonomy-based analysis of pathways and genomes. Nucleic Acids Research, 53(D1), D545-D552. https://doi.org/10.1093/nar/gkae987",
    "Kolberg, L., Raudvere, U., Kuzmin, I., Vilo, J., & Peterson, H. (2020). g:Profiler—Interoperable web service for functional enrichment analysis and gene identifier mapping (2021 update). Nucleic Acids Research, 48(W1), W191-W198. https://doi.org/10.1093/nar/gkaa427",
    "Love, M. I., Huber, W., & Anders, S. (2014). Moderated estimation of fold change and dispersion for RNA-seq data with DESeq2. Genome Biology, 15(12), 550. https://doi.org/10.1186/s13059-014-0550-8",
    "National Center for Biotechnology Information. (2026). BioProject PRJNA1322439. https://www.ncbi.nlm.nih.gov/bioproject/PRJNA1322439",
    "National Center for Biotechnology Information. (2026). GEO series GSE243308. https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE243308",
    "National Center for Biotechnology Information. (2026). Sequence Read Archive study SRP618841. https://www.ncbi.nlm.nih.gov/Traces/study/?acc=SRP618841",
    "The Gene Ontology Consortium. (2021). The Gene Ontology resource: Enriching a GOLD mine. Nucleic Acids Research, 49(D1), D325-D334. https://doi.org/10.1093/nar/gkaa1113",
    "Wirth, R., & Hipp, J. (2000). CRISP-DM: Towards a standard process model for data mining. In Proceedings of the 4th International Conference on the Practical Applications of Knowledge Discovery and Data Mining (pp. 29-39).",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color: str = "666666", size: str = "8") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def citation_text(keys: str) -> str:
    parts = [CITATION_MAP.get(key.strip(), key.strip()) for key in keys.split(",")]
    return "(" + "; ".join(parts) + ")"


def clean_tex(text: str) -> str:
    text = re.sub(r"\\cellcolor\{[^}]*\}", "", text)
    text = re.sub(r"\\codecell\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\inlinecodebox\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\artifact\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\tool\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\textbf\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\textit\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\parencite\{([^}]*)\}", lambda m: citation_text(m.group(1)), text)
    text = text.replace("\\&", "&")
    text = text.replace("\\%", "%")
    text = text.replace("\\_", "_")
    text = text.replace("\\sim", "~")
    text = text.replace("\\textwidth", "")
    text = re.sub(r"\\\((.*?)\\\)", lambda m: m.group(1), text)
    text = re.sub(r"\$([^$]*)\$", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?", "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def parse_table_block(block: str) -> tuple[str, List[List[str]]]:
    caption = clean_tex(re.search(r"\\caption\{(.*?)\}", block, re.S).group(1))
    tabular_match = re.search(
        r"\\begin\{tabularx\}\{.*?\}\{.*?\}(.*?)\\end\{tabularx\}",
        block,
        re.S,
    )
    if tabular_match is None:
        tabular_match = re.search(
            r"\\begin\{tabular\}\{.*?\}(.*?)\\end\{tabular\}",
            block,
            re.S,
        )
    if tabular_match is None:
        raise ValueError("Could not parse table block")
    tabular = tabular_match.group(1)
    tabular_lines = tabular.splitlines()
    while tabular_lines:
        first = tabular_lines[0].strip()
        if not first:
            tabular_lines.pop(0)
            continue
        if "&" not in first and ("p{" in first or "X|" in first or "Y|" in first):
            tabular_lines.pop(0)
            continue
        break
    tabular = "\n".join(tabular_lines)
    tabular = tabular.replace("\r", "")
    tabular = re.sub(r"\\rowcolor\{[^}]*\}", "", tabular)
    tabular = re.sub(r"\\cellcolor\{[^}]*\}", "", tabular)
    tabular = re.sub(r"\\hline", "\\\\HLINE", tabular)
    chunks = [chunk.strip() for chunk in tabular.split("\\HLINE") if chunk.strip()]
    parsed: List[List[str]] = []
    for row in chunks:
        row = row.rstrip("\\").strip()
        if "&" not in row:
            continue
        row = row.replace(r"\&", "__AMPERSAND__")
        cells = [clean_tex(cell).replace("__AMPERSAND__", "&") for cell in row.split("&")]
        parsed.append(cells)
    return caption, parsed


def apply_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str, fill: str | None = None, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(6)
    if fill:
        set_paragraph_shading(p, fill)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1" if level == 1 else "Heading 2"]
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14 if level == 1 else 12)
    r.bold = True


def add_caption(doc: Document, label: str, number: int, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{label} {number}. ")
    r1.bold = True
    r2 = p.add_run(text)
    for r in (r1, r2):
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)


def add_key_table(doc: Document) -> None:
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    headers = [("Member", "Note"),
               ("Nikhi Boggavarapu", "Highlighted text in light purple indicates Nikhi's contribution."),
               ("Sam Kopelev", "Highlighted text in light green indicates Sam's contribution."),
               ("Piter Garcia", "Highlighted text in light red indicates Piter's contribution, including transition wording added for flow.")]
    fills = [None, MEMBER_FILLS["nikhipara"], MEMBER_FILLS["sampara"], MEMBER_FILLS["piterpara"]]
    for i, row in enumerate(headers):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = p.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            if i == 0:
                run.bold = True
            if fills[i]:
                set_cell_shading(cell, fills[i])
    doc.add_paragraph()


def add_table(doc: Document, caption: str, rows: List[List[str]], number: int) -> None:
    add_caption(doc, "Table", number, caption)
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    for i, row in enumerate(rows):
        for j in range(max_cols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            txt = row[j] if j < len(row) else ""
            r = p.add_run(txt)
            r.font.name = "Times New Roman"
            r.font.size = Pt(9)
            if i == 0:
                r.bold = True
                set_cell_shading(cell, "F5F6F7")
            else:
                if j in (0, 1):
                    stage_name = row[0].strip()
                    fill = STAGE_FILLS.get(stage_name)
                    if fill:
                        set_cell_shading(cell, fill)
                    if j == 0:
                        r.bold = True
                elif j == 2:
                    set_cell_shading(cell, "F5F6F7")
                    r.font.name = "Courier New"
                    r.font.size = Pt(8)
    doc.add_paragraph()


def add_figure(doc: Document, image_paths: List[str], caption: str, number: int) -> None:
    width_by_number = {
        2: 6.4,
        6: 6.4,
        7: 6.5,
    }
    single_width = Inches(width_by_number.get(number, 6.0))
    if len(image_paths) == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(str((PAPER_DIR / image_paths[0]).resolve()), width=single_width)
    else:
        table = doc.add_table(rows=1, cols=len(image_paths))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, image in enumerate(image_paths):
            cell = table.cell(0, idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run()
            run.add_picture(str((PAPER_DIR / image).resolve()), width=Inches(3.0))
    add_caption(doc, "Figure", number, caption)


def build_docx() -> None:
    tex = TEX_PATH.read_text()
    body = tex.split(r"\begin{document}", 1)[1].split(r"\end{document}", 1)[0]
    body = re.sub(r"\\maketitle", "", body)
    body = re.sub(r"\\printbibliography", "", body)

    doc = Document()
    apply_doc_defaults(doc)

    for text, size, bold in [
        ("NGS Reanalysis Study on Identification Using Global DEG Discovery of DRG Mouse Dataset", 15, True),
        ("BIOL550 Group Paper Draft 2", 12, False),
        ("Nikhi Boggavarapu | Sam Kopelev | Piter Garcia", 12, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
    doc.add_paragraph()
    add_key_table(doc)

    tokens = re.split(
        r"(\\section\{.*?\}|\\subsection\{.*?\}|\\begin\{nikhipara\}.*?\\end\{nikhipara\}|\\begin\{sampara\}.*?\\end\{sampara\}|\\begin\{piterpara\}.*?\\end\{piterpara\}|\\begin\{figure\}\[H\].*?\\end\{figure\}|\\begin\{table\}\[H\].*?\\end\{table\})",
        body,
        flags=re.S,
    )

    fig_num = 1
    table_num = 1
    first_heading = True
    for tok in tokens:
        tok = tok.strip()
        if not tok:
            continue
        if tok.startswith("\\section{"):
            if not first_heading:
                doc.add_page_break()
            add_heading(doc, clean_tex(re.search(r"\\section\{(.*?)\}", tok, re.S).group(1)), 1)
            first_heading = False
        elif tok.startswith("\\subsection{"):
            add_heading(doc, clean_tex(re.search(r"\\subsection\{(.*?)\}", tok, re.S).group(1)), 2)
        elif tok.startswith("\\begin{nikhipara}"):
            text = clean_tex(re.search(r"\\begin\{nikhipara\}(.*?)\\end\{nikhipara\}", tok, re.S).group(1))
            add_paragraph(doc, text, MEMBER_FILLS["nikhipara"])
        elif tok.startswith("\\begin{sampara}"):
            text = clean_tex(re.search(r"\\begin\{sampara\}(.*?)\\end\{sampara\}", tok, re.S).group(1))
            add_paragraph(doc, text, MEMBER_FILLS["sampara"])
        elif tok.startswith("\\begin{piterpara}"):
            text = clean_tex(re.search(r"\\begin\{piterpara\}(.*?)\\end\{piterpara\}", tok, re.S).group(1))
            add_paragraph(doc, text, MEMBER_FILLS["piterpara"])
        elif tok.startswith("\\begin{figure}"):
            images = re.findall(r"\\includegraphics(?:\[.*?\])?\{(.*?)\}", tok, re.S)
            caption = clean_tex(re.search(r"\\caption\{(.*?)\}", tok, re.S).group(1))
            add_figure(doc, images, caption, fig_num)
            fig_num += 1
        elif tok.startswith("\\begin{table}"):
            caption, rows = parse_table_block(tok)
            add_table(doc, caption, rows, table_num)
            table_num += 1

    doc.add_page_break()
    add_heading(doc, "References", 1)
    for ref in REFERENCE_TEXT:
        add_paragraph(doc, ref)

    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    build_docx()
