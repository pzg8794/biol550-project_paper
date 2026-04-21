#!/usr/bin/env python3
from __future__ import annotations

import re
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

PAPER_DIR = Path(__file__).resolve().parent
TEX_PATH = PAPER_DIR / "materials_methods_piter_draft.tex"
DOCX_PATH = PAPER_DIR / "materials_methods_piter_draft_clean.docx"

STAGE_COLORS = {
    "stagecollect": "365C8D",
    "stageclean": "2E7D6E",
    "stageprep": "C97A2B",
    "stagemine": "7B4FA0",
    "tableheader": "F2F2F2",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color: str = "666666", size: str = "8") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def clean_tex(text: str) -> str:
    text = re.sub(r"\\artifact\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\textbf\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\textit\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\stagecell\{[^}]*\}\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\metriccell\{[^}]*\}\{([^}]*)\}", r"\1", text)
    text = re.sub(r"Figure~\\ref\{[^}]*\}", "the corresponding figure", text)
    text = re.sub(r"Table~\\ref\{[^}]*\}", "the corresponding table", text)
    text = text.replace("\\rightarrow", "->")
    text = text.replace("\\&", "&")
    text = text.replace("\\%", "%")
    text = text.replace("\\_", "_")
    text = text.replace("~", " ")
    text = re.sub(r"\$([^$]*)\$", lambda m: m.group(1).replace("\\rightarrow", "->"), text)
    text = re.sub(r"\\[a-zA-Z]+", "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text)
    text = text.replace("__AMPAMP__", "&&")
    return text.strip()


def parse_table_block(block: str) -> tuple[str, List[List[str]]]:
    caption = clean_tex(re.search(r"\\caption\{(.*?)\}", block, re.S).group(1))
    tabular = re.search(r"\\begin\{tabularx\}\{.*?\}(.*?)\\end\{tabularx\}", block, re.S).group(1)
    tabular = tabular.replace("\r", "")
    tabular = tabular.replace("&&", "__AMPAMP__")
    tabular = re.sub(r"^\{\|.*?\|\}\s*", "", tabular, flags=re.S)
    tabular = re.sub(r"\\rowcolor\{[^}]*\}", "", tabular)
    tabular = re.sub(r"\\Hline[A-Za-z]+", "\\\\HLINE", tabular)
    tabular = re.sub(r"\\hline", "\\\\HLINE", tabular)
    chunks = [chunk.strip() for chunk in tabular.split("\\HLINE") if chunk.strip()]
    parsed = []
    for row in chunks:
        row = row.rstrip("\\").strip()
        cells = [clean_tex(cell.replace("\\\\", "\n")) for cell in row.split("&")]
        parsed.append(cells)
    return caption, parsed


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def apply_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(6)


def add_caption(doc: Document, label: str, number: int, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(f"{label} {number}. ")
    r1.bold = True
    r2 = p.add_run(text)
    for r in (r1, r2):
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)


def add_table(doc: Document, caption: str, rows: List[List[str]], number: int) -> None:
    add_caption(doc, "Table", number, caption)
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    for i, row in enumerate(rows):
        for j in range(max_cols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[j] if j < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(1)
            for line_idx, line in enumerate(text.split("\n")):
                if line_idx > 0:
                    p.add_run().add_break(WD_BREAK.LINE)
                run = p.add_run(line)
                run.font.name = "Times New Roman"
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True
            if i == 0:
                set_cell_shading(cell, STAGE_COLORS["tableheader"])
            elif j == 0:
                set_cell_shading(cell, "EEEAF7")
    doc.add_paragraph()


def add_figure(doc: Document, image_rel: str, caption: str, number: int) -> None:
    image_path = (PAPER_DIR / image_rel).resolve()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.4))
    add_caption(doc, "Figure", number, caption)


def build_docx() -> None:
    tex = TEX_PATH.read_text()
    body = tex.split(r"\begin{document}", 1)[1].split(r"\end{document}", 1)[0]
    body = re.sub(r"\\maketitle", "", body)

    doc = Document()
    apply_doc_defaults(doc)

    # Title block
    for text, size, bold in [
        ("Materials and Methods Draft", 16, True),
        ("Mouse DRG RNA-seq Follow-up Project", 14, False),
        ("Piter Garcia", 12, False),
        ("Draft prepared for BIOL550", 12, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold

    tokens = re.split(r"(\\section\{.*?\}|\\subsection\{.*?\}|\\begin\{figure\}\[H\].*?\\end\{figure\}|\\begin\{table\}\[H\].*?\\end\{table\})", body, flags=re.S)
    fig_num = 1
    table_num = 1
    for tok in tokens:
        tok = tok.strip()
        if not tok:
            continue
        if tok.startswith("\\section{"):
            text = clean_tex(re.search(r"\\section\{(.*?)\}", tok, re.S).group(1))
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 1"]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            r = p.add_run(text)
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
            r.bold = True
        elif tok.startswith("\\subsection{"):
            text = clean_tex(re.search(r"\\subsection\{(.*?)\}", tok, re.S).group(1))
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 2"]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            r = p.add_run(text)
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.bold = True
        elif tok.startswith("\\begin{figure}"):
            image_match = re.search(r"\\includegraphics(?:\[.*?\])?\{(.*?)\}", tok, re.S)
            caption = clean_tex(re.search(r"\\caption\{(.*?)\}", tok, re.S).group(1))
            if image_match is not None:
                image = image_match.group(1)
            else:
                label_match = re.search(r"\\label\{(.*?)\}", tok, re.S)
                label = label_match.group(1) if label_match else ""
                image = "assets_methods/overview_pipeline_stage.png" if label == "fig:overall_pipeline" else None
            if image is not None:
                add_figure(doc, image, caption, fig_num)
            else:
                add_caption(doc, "Figure", fig_num, caption)
            fig_num += 1
        elif tok.startswith("\\begin{table}"):
            caption, rows = parse_table_block(tok)
            add_table(doc, caption, rows, table_num)
            table_num += 1
        else:
            paragraphs = [clean_tex(p) for p in re.split(r"\n\s*\n", tok) if clean_tex(p)]
            for para in paragraphs:
                add_paragraph(doc, para)

    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    build_docx()
